from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self):
        self.document = Document()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """Настраивает стили документа"""
        # Стиль для заголовков
        heading_style = self.document.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_after = Pt(12)
        
        # Стиль для подзаголовков
        subheading_style = self.document.styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_after = Pt(8)
        
        # Стиль для обычного текста
        normal_style = self.document.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
    
    def add_title(self, title):
        """Добавляет заголовок документа"""
        title_paragraph = self.document.add_paragraph(title)
        title_paragraph.style = self.document.styles['CustomHeading']
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section_heading(self, heading):
        """Добавляет заголовок раздела"""
        heading_paragraph = self.document.add_paragraph(heading)
        heading_paragraph.style = self.document.styles['CustomSubheading']
    
    def add_paragraph(self, text):
        """Добавляет параграф текста"""
        paragraph = self.document.add_paragraph(text)
        paragraph.style = self.document.styles['CustomNormal']
    
    def add_dialog_entry(self, speaker, message, timestamp):
        """Добавляет запись диалога"""
        # Форматируем время
        time_str = timestamp.strftime("%H:%M:%S")
        
        # Добавляем запись диалога
        dialog_text = f"[{time_str}] {speaker}: {message}"
        self.add_paragraph(dialog_text)
    
    def generate_report(self, user_id, conversation_history, analytics_report):
        """Генерирует полный отчет"""
        # Создаем новый документ
        self.document = Document()
        self.setup_document_styles()
        
        # Добавляем заголовок
        current_time = datetime.now()
        title = f"Отчет по собеседованию\n{current_time.strftime('%d.%m.%Y %H:%M')}"
        self.add_title(title)
        
        # Добавляем информацию о пользователе
        self.add_section_heading("Информация о кандидате")
        self.add_paragraph(f"ID пользователя: {user_id}")
        self.add_paragraph(f"Дата собеседования: {current_time.strftime('%d.%m.%Y')}")
        self.add_paragraph(f"Время начала: {conversation_history[0]['timestamp'].strftime('%H:%M')}")
        self.add_paragraph(f"Время завершения: {current_time.strftime('%H:%M')}")
        
        # Добавляем раздел с диалогом
        self.add_section_heading("Диалог собеседования")
        
        for msg in conversation_history:
            speaker = "Рекрутер" if msg["is_bot"] else "Кандидат"
            self.add_dialog_entry(speaker, msg["text"], msg["timestamp"])
        
        # Добавляем аналитический отчет
        self.add_section_heading("Аналитический отчет")
        
        # Разбиваем отчет на параграфы
        report_lines = analytics_report.split('\n')
        current_paragraph = ""
        
        for line in report_lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    self.add_paragraph(current_paragraph)
                    current_paragraph = ""
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
                # Это заголовок раздела
                if current_paragraph:
                    self.add_paragraph(current_paragraph)
                    current_paragraph = ""
                self.add_section_heading(line)
            else:
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        
        # Добавляем последний параграф, если есть
        if current_paragraph:
            self.add_paragraph(current_paragraph)
    
    def save_document(self, user_id):
        """Сохраняет документ в папку dialogs"""
        # Создаем папку, если её нет
        os.makedirs("dialogs", exist_ok=True)
        
        # Формируем имя файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"dialogs/interview_report_{user_id}_{timestamp}.docx"
        
        # Сохраняем документ
        self.document.save(filename)
        return filename

